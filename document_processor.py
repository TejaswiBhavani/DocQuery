import PyPDF2
import re
import email
import os
from typing import List, Optional
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    
try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False

class DocumentProcessor:
    """Handles document processing and text extraction for multiple formats."""
    
    def detect_file_type(self, file_path: str) -> str:
        """Detect file type based on content and extension."""
        # Fallback to extension-based detection
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.pdf':
            return 'pdf'
        elif ext in ['.docx', '.doc']:
            return 'docx'
        elif ext in ['.eml', '.msg', '.txt']:
            return 'email'
        else:
            return 'unknown'
    
    def extract_text(self, file_path: str) -> str:
        """Extract text content from various document formats."""
        file_type = self.detect_file_type(file_path)
        
        try:
            if file_type == 'pdf':
                return self._extract_pdf_text(file_path)
            elif file_type == 'docx':
                return self._extract_docx_text(file_path)
            elif file_type == 'email':
                return self._extract_email_text(file_path)
            else:
                # Try to read as plain text
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    return self._clean_text(file.read())
                    
        except Exception as e:
            raise Exception(f"Error extracting text from {file_type} file: {str(e)}")
    
    def _extract_pdf_text(self, pdf_path: str) -> str:
        """Extract text content from a PDF file."""
        try:
            text_content = ""
            
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Check if PDF is encrypted
                if pdf_reader.is_encrypted:
                    raise Exception("PDF is encrypted and cannot be processed")
                
                # Extract text from all pages
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
                
                if not text_content.strip():
                    raise Exception("No text content found in PDF")
                
                return self._clean_text(text_content)
                
        except PyPDF2.errors.PdfReadError as e:
            raise Exception(f"Error reading PDF file: {str(e)}")
        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {str(e)}")
    
    def _extract_docx_text(self, docx_path: str) -> str:
        """Extract text content from a Word document."""
        if not DOCX_AVAILABLE:
            raise Exception("Word document processing not available. Please install python-docx.")
        
        try:
            doc = Document(docx_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return self._clean_text(text)
            
        except Exception as e:
            raise Exception(f"Error extracting text from Word document: {str(e)}")
    
    def _extract_email_text(self, email_path: str) -> str:
        """Extract text content from an email file."""
        try:
            with open(email_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
            
            # Try to parse as email
            try:
                msg = email.message_from_string(content)
                
                # Extract email metadata
                subject = msg.get('Subject', '')
                sender = msg.get('From', '')
                recipient = msg.get('To', '')
                date = msg.get('Date', '')
                
                # Extract body
                body = ""
                if msg.is_multipart():
                    for part in msg.walk():
                        if part.get_content_type() == "text/plain":
                            payload = part.get_payload(decode=True)
                            if payload:
                                body += payload.decode('utf-8', errors='ignore')
                else:
                    payload = msg.get_payload(decode=True)
                    if payload:
                        body = payload.decode('utf-8', errors='ignore')
                    else:
                        body = str(msg.get_payload())
                
                # Combine all content
                email_text = f"""Subject: {subject}
From: {sender}
To: {recipient}
Date: {date}

Body:
{body}"""
                
                return self._clean_text(email_text)
                
            except Exception:
                # If email parsing fails, return as plain text
                return self._clean_text(content)
                
        except Exception as e:
            raise Exception(f"Error extracting text from email: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """
        Clean extracted text by removing extra whitespace and formatting issues.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove page numbers and headers/footers (basic patterns)
        text = re.sub(r'\n\s*\d+\s*\n', '\n', text)
        text = re.sub(r'\n\s*Page\s+\d+.*?\n', '\n', text, flags=re.IGNORECASE)
        
        # Fix common OCR issues
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)  # Add space between joined words
        
        # Remove extra newlines
        text = re.sub(r'\n+', '\n', text)
        
        return text.strip()
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """
        Split text into overlapping chunks for better semantic search.
        
        Args:
            text: Text to chunk
            chunk_size: Maximum size of each chunk
            overlap: Number of characters to overlap between chunks
            
        Returns:
            List of text chunks
        """
        if not text or not text.strip():
            return []
        
        # First, try to split by paragraphs
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        
        chunks = []
        current_chunk = ""
        
        for paragraph in paragraphs:
            # If adding this paragraph would exceed chunk size
            if len(current_chunk) + len(paragraph) > chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                
                # Start new chunk with overlap from previous chunk
                if overlap > 0 and len(current_chunk) > overlap:
                    current_chunk = current_chunk[-overlap:] + " " + paragraph
                else:
                    current_chunk = paragraph
            else:
                current_chunk += " " + paragraph if current_chunk else paragraph
        
        # Add the last chunk
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        # If no chunks were created (very long paragraphs), fall back to fixed-size chunking
        if not chunks:
            chunks = self._fixed_size_chunking(text, chunk_size, overlap)
        
        return [chunk for chunk in chunks if len(chunk.strip()) > 50]  # Filter out very short chunks
    
    def _fixed_size_chunking(self, text: str, chunk_size: int, overlap: int) -> List[str]:
        """
        Fallback method for fixed-size text chunking.
        
        Args:
            text: Text to chunk
            chunk_size: Size of each chunk
            overlap: Overlap between chunks
            
        Returns:
            List of text chunks
        """
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to end at a sentence boundary
            if end < len(text):
                # Look for sentence endings within the last 100 characters
                sentence_end = text.rfind('.', start, end)
                if sentence_end > start + chunk_size - 100:
                    end = sentence_end + 1
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap if overlap > 0 else end
        
        return chunks
